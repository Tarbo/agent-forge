"""
Word document export tool.

Uses python-docx to create formatted .docx files with dynamic formatting support.
"""
from pathlib import Path
from langchain_core.tools import tool
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.utils.file_manager import get_full_path
from src.utils.logger import logger


# Property registry - easily extensible as python-docx adds more features
FONT_PROPERTIES = {
    "name": str,          # Font name (e.g., "Arial", "Calibri")
    "size": Pt,           # Font size (needs Pt conversion)
    "bold": bool,         # Bold text
    "italic": bool,       # Italic text
    "underline": bool,    # Underline text
}

PARAGRAPH_PROPERTIES = {
    "alignment": int,     # Alignment (use WD_ALIGN_PARAGRAPH constants)
    "line_spacing": float,  # Line spacing multiplier
}

TITLE_PROPERTIES = {
    "alignment": int,     # Title alignment
}

# Default formatting values
DEFAULT_FORMATTING = {
    "name": "Calibri",
    "size": 11,
}

DEFAULT_TITLE_FORMATTING = {
    "alignment": WD_ALIGN_PARAGRAPH.LEFT,
}


@tool
def export_to_word(text: str, custom_name: str = None, formatting: dict = None) -> str:
    """
    Export text content as a Microsoft Word document (.docx) with dynamic formatting.
    
    This tool creates a formatted Word document with the provided text.
    Accepts any formatting options via dict - unsupported options are gracefully ignored.
    
    Args:
        text: The text content to export to Word
        custom_name: Optional custom filename (without extension)
        formatting: Dict of formatting options (see FONT_PROPERTIES and PARAGRAPH_PROPERTIES)
    
    Supported formatting (extensible via registries):
        Font: name, size, bold, italic, underline
        Paragraph: alignment, line_spacing
        Title: title_alignment (prefix with 'title_' for title-specific props)
    
    Returns:
        str: Full path to the created Word document
        
    Example:
        export_to_word("My Title\n\nContent here", "my_report", 
                       formatting={"name": "Arial", "size": 14, "bold": True,
                                   "title_alignment": WD_ALIGN_PARAGRAPH.CENTER})
        # Returns: "/path/to/exports/my_report_2024-10-23_14-30-22.docx"
    """
    try:
        # Generate file path
        file_path = get_full_path("word", custom_name)
        
        # Merge user formatting with defaults
        format_opts = {**DEFAULT_FORMATTING, **(formatting or {})}
        
        # Create Document
        doc = Document()
        
        # Add title (first line if available)
        lines = text.strip().split('\n')
        if lines:
            # First line as title
            title = doc.add_heading(lines[0][:100], level=1)
            
            # Apply title formatting dynamically from registry
            title_opts = {**DEFAULT_TITLE_FORMATTING, **{k.replace("title_", ""): v for k, v in format_opts.items() if k.startswith("title_")}}
            for prop_name, prop_type in TITLE_PROPERTIES.items():
                if prop_name in title_opts:
                    try:
                        setattr(title, prop_name, title_opts[prop_name])
                        logger.debug(f"Applied title.{prop_name} = {title_opts[prop_name]}")
                    except Exception as e:
                        logger.warning(f"Failed to apply title.{prop_name}: {e}")
            
            # Rest as content
            remaining_text = '\n'.join(lines[1:]) if len(lines) > 1 else ""
        else:
            remaining_text = text
        
        # Add paragraphs with dynamic formatting
        if remaining_text.strip():
            paragraphs = remaining_text.split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    # Create empty paragraph first
                    para = doc.add_paragraph()
                    
                    # Add run with text (this allows us to format it)
                    run = para.add_run(para_text.strip())
                    
                    # Apply font properties dynamically from registry
                    for prop_name, prop_type in FONT_PROPERTIES.items():
                        if prop_name in format_opts:
                            try:
                                value = format_opts[prop_name]
                                # Convert size to Pt if needed
                                if prop_name == "size" and prop_type == Pt:
                                    value = Pt(value)
                                setattr(run.font, prop_name, value)
                                logger.debug(f"Applied font.{prop_name} = {value}")
                            except Exception as e:
                                logger.warning(f"Failed to apply font.{prop_name}: {e}")
                    
                    # Apply paragraph properties dynamically from registry
                    for prop_name, prop_type in PARAGRAPH_PROPERTIES.items():
                        if prop_name in format_opts:
                            try:
                                value = format_opts[prop_name]
                                setattr(para, prop_name, value)
                                logger.debug(f"Applied paragraph.{prop_name} = {value}")
                            except Exception as e:
                                logger.warning(f"Failed to apply paragraph.{prop_name}: {e}")
        
        # Save document
        doc.save(str(file_path))
        
        logger.info(f"Word document created: {file_path}")
        return str(file_path)
        
    except Exception as e:
        error_msg = f"Failed to create Word document: {e}"
        logger.error(error_msg)
        raise RuntimeError(error_msg)

